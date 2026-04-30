import docx

out_path = r"C:\Users\Priya\.gemini\antigravity\scratch\tata_iq_eda\Task3_Business_Summary_Report.docx"

doc = docx.Document()

doc.add_heading("Business Summary Report: Predictive Insights for Collections Strategy", 0)

# Section 1
doc.add_heading("1. Summary of Predictive Insights", level=1)
doc.add_paragraph("The predictive model analyzed 500 customer records to identify high-risk segments and key predictors of delinquency. The results show that behavioral financial indicators heavily outweigh demographic data in forecasting default risk. Customers in the highest risk segment typically exhibit a combination of maxed-out credit limits and recent payment friction.")

doc.add_paragraph("Top 3 Risk Factors for Delinquency:", style='List Bullet')
doc.add_paragraph("High Credit Utilization: Customers maximizing their available credit limits are the most vulnerable to cash flow disruptions and subsequent delinquency.", style='List Bullet')
doc.add_paragraph("Recent Missed Payments: A history of missed or late payments in the trailing 6 months (Months 1-6) is the strongest sequential indicator of escalating financial distress.", style='List Bullet')
doc.add_paragraph("High Debt-to-Income (DTI) Ratio: Customers with excessive debt relative to their income have minimal financial flexibility, increasing default probability.", style='List Bullet')

# Section 2
doc.add_heading("2. Recommendation Framework", level=1)
p = doc.add_paragraph()
p.add_run("Restated Insight: ").bold = True
p.add_run("High credit utilization combined with recent missed payments significantly increases the likelihood of long-term delinquency.")

p2 = doc.add_paragraph()
p2.add_run("Proposed Recommendation: ").bold = True
p2.add_run("Implement an automated 'Early Intervention Protocol' for customers exceeding 80% credit utilization who miss a single payment.")

doc.add_paragraph("Specific: Trigger a personalized financial wellness email and SMS offering flexible repayment options to any customer who exceeds 80% credit utilization and misses one monthly payment.", style='List Bullet 2')
doc.add_paragraph("Measurable: Aim to reduce the progression of first-time missed payments to 60-day delinquency by 15%.", style='List Bullet 2')
doc.add_paragraph("Actionable: The Analytics team will configure the predictive model to flag these specific users weekly, and the Collections team will automate the outreach workflow in their CRM.", style='List Bullet 2')
doc.add_paragraph("Relevant: Proactive intervention reduces ultimate default rates, aligning perfectly with Geldium’s goal of lowering financial loss while preserving customer relationships.", style='List Bullet 2')
doc.add_paragraph("Time-bound: Deploy the pilot program within 30 days and measure its impact over the subsequent quarter (90 days).", style='List Bullet 2')

p3 = doc.add_paragraph()
p3.add_run("Justification and Business Rationale: ").bold = True
p3.add_run("Waiting until a customer misses multiple payments drastically reduces the chance of full recovery. By using the predictive model to intervene early—specifically when a customer shows the dual warning signs of maxed-out credit and a single missed payment—Geldium can proactively offer assistance. This shifts the collections strategy from punitive to supportive, ultimately improving debt recovery rates and customer retention.")

# Section 3
doc.add_heading("3. Ethical and Responsible AI Considerations", level=1)
doc.add_paragraph("Fairness Risks and Mitigations:", style='Intense Quote')

p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Demographic Bias: ").bold = True
p4.add_run("The model could inadvertently learn to penalize younger demographics or specific regions if historical data shows higher default rates there, leading to discriminatory collection practices. ")
p4.add_run("Mitigation: ").italic = True
p4.add_run("We will actively exclude explicit demographic variables (Age, Location) from the final decision-making logic and rely strictly on financial behavior. We will also run Disparate Impact checks to ensure flag rates are equitable across groups.")

p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Data Quality Bias: ").bold = True
p5.add_run("Missing income data (~7.8% of the dataset) was imputed using median values. If higher-income individuals are more likely to withhold income data, imputing the median might incorrectly assess their Debt-to-Income ratio, flagging them inaccurately. ")
p5.add_run("Mitigation: ").italic = True
p5.add_run("We will implement a 'human-in-the-loop' review for high-value accounts with imputed data before initiating aggressive collection actions.")

doc.add_paragraph("Explainability and Transparency:", style='Intense Quote')
doc.add_paragraph("The Random Forest model was selected specifically for its explainability. Instead of a 'black box' prediction, we can extract the exact features driving a risk score. When communicating with stakeholders or customers, we can plainly state: 'This account was flagged because credit utilization reached 90% and the last payment was 15 days late,' ensuring transparency and providing actionable feedback rather than abstract AI probabilities. This approach promotes responsible and transparent AI use.")

doc.save(out_path)
print("Saved fresh docx to:", out_path)
